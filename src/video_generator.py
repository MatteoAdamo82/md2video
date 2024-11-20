import os
from docx import Document
from datetime import datetime
import logging
from typing import Dict, List, Optional, Callable
from dotenv import load_dotenv
from .blog_parser import BlogParser
from .video_maker import VideoMaker
import sys
from pathlib import Path

class VideoGenerator:
    def __init__(self):
        load_dotenv()
        self.content_dir = os.getenv('HUGO_CONTENT_DIR', 'content/posts')
        self.num_posts = int(os.getenv('NUM_POSTS', '5'))
        self.output_dir = os.getenv('SCRIPT_OUTPUT_DIR', 'video_scripts')
        self.video_output_dir = os.getenv('VIDEO_OUTPUT_DIR', 'video_output')

        # Messaggi predefiniti
        self.intro_text = os.getenv('VIDEO_INTRO_TEXT', 'Ciao a tutti e bentornati sul canale!')
        self.outro_text = os.getenv('VIDEO_OUTRO_TEXT', 'Grazie per aver guardato questo video!')

        # Configura logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)

        # Crea directory se non esistono
        for directory in [self.output_dir, self.video_output_dir]:
            Path(directory).mkdir(parents=True, exist_ok=True)

    def generate_scripts(self, message_callback=None, progress_callback=None):
            """Genera solo gli script dai post"""
            try:
                self.log_message(message_callback, "üì• Fetching recent posts...")
                parser = BlogParser(self.content_dir)
                posts = parser.fetch_posts(self.num_posts)

                if not posts:
                    self.log_message(message_callback, "‚ùå No posts found.")
                    return []

                results = []
                total_posts = len(posts)

                for i, post in enumerate(posts, 1):
                    try:
                        self.update_progress(progress_callback, (i * 100) // total_posts,
                                           f"Creating script {i} of {total_posts}")
                        script_file, _ = self.create_script(post)
                        results.append({
                            'title': post['title'],
                            'script_file': script_file,
                            'url': post['url']
                        })
                        self.log_message(message_callback, f"‚úÖ Created script for: {post['title']}")
                    except Exception as e:
                        self.log_message(message_callback, f"‚ùå Error: {str(e)}")

                return results
            except Exception as e:
                raise Exception(f"Error generating scripts: {str(e)}")

    def list_available_scripts(self):
        """Lista gli script disponibili"""
        try:
            return list(Path(self.output_dir).glob('*.docx'))
        except Exception as e:
            raise Exception(f"Error listing scripts: {str(e)}")

    def generate_video_from_script(self, script_path: str, message_callback=None,
                                 progress_callback=None):
        """Genera un video da uno script esistente"""
        try:
            doc = Document(script_path)
            # Qui andrebbe la logica di conversione da docx a script_content
            # Per ora restituiamo un errore
            raise NotImplementedError("Video generation from script not yet implemented")
        except Exception as e:
            raise Exception(f"Error generating video: {str(e)}")

    def generate_video_from_script(self, script_path: str, message_callback=None,
                                     progress_callback=None):
            """Genera un video da uno script esistente"""
            try:
                self.log_message(message_callback, "üìÑ Lettura script...")
                doc = Document(script_path)

                script_content = {
                    'title': doc.paragraphs[0].text.replace('Script Video: ', ''),
                    'sections': []
                }

                # Estrai le sezioni
                current_section = None
                for para in doc.paragraphs[3:]:  # Salta intestazione e metadati
                    if para.style.name.startswith('Heading'):
                        if current_section:
                            script_content['sections'].append(current_section)
                        current_section = {
                            'level': int(para.style.name[-1]),
                            'title': para.text,
                            'content': []
                        }
                    elif current_section and para.text:
                        current_section['content'].append(para.text)

                if current_section:
                    script_content['sections'].append(current_section)

                self.update_progress(progress_callback, 30, "Script analizzato")

                # Genera il video
                video_maker = VideoMaker(self.video_output_dir)
                video_file = video_maker.create_video(script_content)

                self.update_progress(progress_callback, 100, "Video completato")
                return video_file

            except Exception as e:
                raise Exception(f"Error generating video: {str(e)}")
    def create_script(self, post: Dict) -> tuple[str, Dict]:
        """Crea uno script formattato dal post"""
        try:
            doc = Document()
            script_content = {
                'title': post['title'],
                'url': post['url'],
                'date': post['date'],
                'sections': []
            }

            # Intestazione del documento
            doc.add_heading(f'Script Video: {post["title"]}', 0)
            doc.add_paragraph(f'Post originale: {post["url"]}')
            doc.add_paragraph(f'Data post: {post["date"]}')

            # Intro
            intro_section = {
                'level': 1,
                'title': 'Introduzione',
                'content': [self.intro_text, f'Oggi parleremo di {post["title"]}']
            }
            script_content['sections'].append(intro_section)

            doc.add_heading('üé¨ Introduzione', 1)
            doc.add_paragraph(self.intro_text)
            doc.add_paragraph(f'Oggi parleremo di {post["title"]}')

            # Contenuto principale
            doc.add_heading('üìù Contenuto', 1)
            if 'sections' in post:
                script_content['sections'].extend(post['sections'])

                for section in post['sections']:
                    if section['title']:
                        doc.add_heading(section['title'], section['level'])
                    if section['content']:
                        for para in section['content']:
                            doc.add_paragraph(para)

            # Outro
            outro_section = {
                'level': 1,
                'title': 'Conclusione',
                'content': [self.outro_text]
            }
            script_content['sections'].append(outro_section)

            doc.add_heading('üé¨ Conclusione', 1)
            doc.add_paragraph(self.outro_text)

            # Salva il documento
            filename = os.path.join(
                self.output_dir,
                f"script_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{post['title'][:30]}.docx"
            )
            doc.save(filename)

            self.logger.info(f"Created script: {filename}")
            return filename, script_content

        except Exception as e:
            error_msg = f"Error creating script for {post['title']}: {str(e)}"
            self.logger.error(error_msg)
            raise Exception(error_msg)

    def update_progress(self, progress_callback: Optional[Callable] = None,
                       progress: float = 0, status: str = ""):
        """Utility method per aggiornare il progresso in modo sicuro"""
        if progress_callback:
            try:
                progress_callback({
                    "value": progress,
                    "status": status
                })
            except Exception as e:
                self.logger.error(f"Error updating progress: {str(e)}")

    def log_message(self, message_callback: Optional[Callable] = None, message: str = ""):
        """Utility method per loggare messaggi in modo sicuro"""
        if message_callback:
            try:
                message_callback(message)
            except Exception as e:
                self.logger.error(f"Error logging message: {str(e)}")

    def process_post(self, post: Dict, message_callback: Optional[Callable] = None,
                    progress_callback: Optional[Callable] = None) -> Dict:
        """Processa un singolo post"""
        try:
            # Inizio processo
            self.update_progress(progress_callback, 0, f"Processing: {post['title']}")
            self.log_message(message_callback, f"Starting processing of: {post['title']}")

            # Crea lo script
            self.log_message(message_callback, f"Creating script for: {post['title']}")
            script_file, script_content = self.create_script(post)
            self.update_progress(progress_callback, 30, f"Script created for: {post['title']}")

            # Crea il video
            self.log_message(message_callback, f"Creating video for: {post['title']}")
            video_maker = VideoMaker(self.video_output_dir)
            video_file = video_maker.create_video(script_content)
            self.update_progress(progress_callback, 100, f"Video created for: {post['title']}")

            return {
                'title': post['title'],
                'script_file': script_file,
                'video_file': video_file,
                'url': post['url']
            }

        except Exception as e:
            error_msg = f"Error processing post {post['title']}: {str(e)}"
            self.logger.error(error_msg)
            self.log_message(message_callback, f"‚ùå {error_msg}")
            raise

    def process_recent_posts(self, message_callback: Optional[Callable] = None,
                           progress_callback: Optional[Callable] = None) -> List[Dict]:
        """Processa i post pi√π recenti"""
        try:
            self.log_message(message_callback, "üì• Fetching recent posts...")

            parser = BlogParser(self.content_dir)
            posts = parser.fetch_posts(self.num_posts)

            if not posts:
                self.log_message(message_callback, "‚ùå No posts found.")
                return []

            processed_items = []
            total_posts = len(posts)

            for i, post in enumerate(posts, 1):
                try:
                    # Calcola e aggiorna il progresso generale
                    overall_progress = ((i - 1) * 100) // total_posts
                    self.update_progress(
                        progress_callback,
                        overall_progress,
                        f"Processing post {i} of {total_posts}"
                    )

                    # Processa il post
                    result = self.process_post(post, message_callback, progress_callback)
                    if result:
                        processed_items.append(result)
                        self.log_message(
                            message_callback,
                            f"‚úÖ Successfully processed: {post['title']}"
                        )

                except Exception as e:
                    self.logger.error(f"Error processing post {post['title']}: {str(e)}")
                    self.log_message(
                        message_callback,
                        f"‚ö†Ô∏è Error processing post {post['title']}: {str(e)}"
                    )
                    continue

            # Completamento
            self.update_progress(progress_callback, 100, "Processing complete")
            return processed_items

        except Exception as e:
            error_msg = f"Critical error: {str(e)}"
            self.logger.error(error_msg)
            self.log_message(message_callback, f"‚ùå {error_msg}")
            raise

def main():
    try:
        generator = VideoGenerator()
        print("üöÄ Starting video generation...")

        def progress_callback(info: Dict):
            progress = info.get('value', 0)
            status = info.get('status', '')
            print(f"Progress: {progress}% - {status}")

        items = generator.process_recent_posts(
            message_callback=print,
            progress_callback=progress_callback
        )

        if items:
            print("\n‚úÖ Created files:")
            for item in items:
                print(f"\nüìÑ {item['title']}")
                print(f"   üìù Script: {item['script_file']}")
                print(f"   üé• Video: {item['video_file']}")
                print(f"   üîó URL: {item['url']}")
        else:
            print("\n‚ùå No files created. Check logs for details.")

    except Exception as e:
        print(f"\n‚ùå Execution error: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()